"""
Text extraction for supported document types.
"""
import re
from pathlib import Path

from docx import Document as DocxDocument
from PyPDF2 import PdfReader


def extract_text(file_path: str, file_ext: str) -> str:
    if file_ext == ".pdf":
        return _extract_pdf(file_path)
    elif file_ext == ".docx":
        return _extract_docx(file_path)
    elif file_ext == ".txt":
        return _extract_txt(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {file_ext}")


def _extract_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs)


def _extract_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def clean_text(text: str) -> str:
    """Normalize whitespace and strip control characters."""
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> list:
    """
    Simple sliding-window chunker over characters, trying to break on
    paragraph/sentence boundaries where possible.
    """
    if chunk_overlap >= chunk_size:
        chunk_overlap = max(0, chunk_size // 4)

    chunks = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = min(start + chunk_size, text_len)

        if end < text_len:
            boundary = text.rfind("\n\n", start, end)
            if boundary == -1:
                boundary = text.rfind(". ", start, end)
            if boundary != -1 and boundary > start + (chunk_size // 2):
                end = boundary + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= text_len:
            break
        start = max(end - chunk_overlap, start + 1)

    return chunks
